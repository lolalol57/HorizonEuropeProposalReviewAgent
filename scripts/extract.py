"""Steps 3 & 5 — Document Extraction and Figure Extraction.

Usage:
    python3 scripts/extract.py <run-id>

Produces (in workspace/<run-id>/internal/):
    document.json    page text + heading structure + full_text
    table-map.json   detected tables (headers, rows) + detected "Table N" labels
    figure-map.json  extracted figure PNGs + detected "Figure N" captions

Figure PNGs are written to workspace/<run-id>/figures/. This step does NOT
evaluate quality — it only determines *what exists and where*. The host model
performs the visual review by opening the extracted PNGs later.

`render_page(run_id, page_no)` is exposed for on-demand page rendering when a
figure could not be extracted directly (never bulk-render every page).
"""
import os
import sys

import fitz  # PyMuPDF

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import common  # noqa: E402

MIN_FIG_PX = 90  # ignore icons/logos smaller than this on a side


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------
def _caption_lines(page, regex):
    """Return [(number:int, title:str, y_top:float)] for lines matching regex."""
    out = []
    data = page.get_text("dict")
    for block in data.get("blocks", []):
        for line in block.get("lines", []):
            text = "".join(span.get("text", "") for span in line.get("spans", [])).strip()
            m = regex.match(text)
            if m:
                out.append((int(m.group(1)), (m.group(2) or "").strip(),
                            float(line["bbox"][1])))
    return out


def _body_font_size(doc):
    """Most common rounded span size = body text size (heading heuristic base)."""
    counts = {}
    for page in doc:
        for block in page.get_text("dict").get("blocks", []):
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    sz = round(span.get("size", 0))
                    txt = span.get("text", "").strip()
                    if txt:
                        counts[sz] = counts.get(sz, 0) + len(txt)
    if not counts:
        return 10.0
    return float(max(counts, key=counts.get))


def _headings(page, body_size):
    """Short, larger-than-body lines treated as candidate headings."""
    out = []
    for block in page.get_text("dict").get("blocks", []):
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            text = "".join(s.get("text", "") for s in spans).strip()
            if not text or len(text) > 140:
                continue
            max_sz = max((s.get("size", 0) for s in spans), default=0)
            bold = any("bold" in (s.get("font", "").lower()) for s in spans)
            if max_sz >= body_size + 1.4 or (bold and max_sz >= body_size + 0.5
                                             and len(text) < 90):
                out.append({"text": text, "size": round(max_sz, 1), "bold": bold})
    return out


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------
def extract_pdf(run_id, path):
    doc = fitz.open(path)
    body_size = _body_font_size(doc)

    pages = []
    full_text_parts = []
    for i in range(doc.page_count):
        page = doc[i]
        text = page.get_text("text")
        full_text_parts.append(text)
        pages.append({
            "page": i + 1,
            "text": text,
            "headings": _headings(page, body_size),
        })

    document = {
        "run_id": run_id,
        "source": os.path.basename(path),
        "type": "pdf",
        "page_count": doc.page_count,
        "body_font_size": body_size,
        "pages": pages,
        "full_text": "\n".join(full_text_parts),
    }
    common.write_json(common.internal_path(run_id, "document.json"), document)

    table_map = _extract_pdf_tables(doc)
    common.write_json(common.internal_path(run_id, "table-map.json"), table_map)

    figure_map = _extract_pdf_figures(doc, run_id)
    common.write_json(common.internal_path(run_id, "figure-map.json"), figure_map)

    doc.close()
    return document, table_map, figure_map


def _extract_pdf_tables(doc):
    tables = []
    labels = []
    tid = 0
    for i in range(doc.page_count):
        page = doc[i]
        pno = i + 1
        caps = _caption_lines(page, common.TABLE_CAPTION_RE)
        for num, title, y0 in caps:
            labels.append({"detected_number": num, "detected_title": title,
                           "page": pno})
        try:
            found = page.find_tables()
        except Exception:
            found = None
        if not found:
            continue
        for t in getattr(found, "tables", []):
            tid += 1
            try:
                rows = t.extract()
            except Exception:
                rows = []
            headers = []
            try:
                if t.header is not None:
                    headers = [h for h in t.header.names if h]
            except Exception:
                headers = []
            if not headers and rows:
                headers = [str(c) for c in rows[0] if c]
            top = t.bbox[1]
            caption = _nearest_caption(caps, top, prefer_above=True)
            tables.append({
                "id": "TAB-{:02d}".format(tid),
                "detected_number": caption[0] if caption else None,
                "detected_title": caption[1] if caption else None,
                "page": pno,
                "column_headers": headers,
                "row_count": len(rows),
                "sample_rows": rows[:3],
            })
    return {"tables": tables, "detected_table_labels": labels}


def _extract_pdf_figures(doc, run_id):
    fig_dir = common.subdir(run_id, "figures")
    figures = []
    unmatched_images = []
    seen = set()
    fid = 0

    # Collect every figure caption up-front (some figures are vector art with no
    # embedded raster image; those still deserve an entry so the model can render
    # the page on demand).
    all_caps = {}  # page -> [(num, title, y0)]
    for i in range(doc.page_count):
        all_caps[i + 1] = _caption_lines(doc[i], common.FIGURE_CAPTION_RE)

    matched_caption_keys = set()

    for i in range(doc.page_count):
        page = doc[i]
        pno = i + 1
        caps = all_caps[pno]
        for img in page.get_images(full=True):
            xref = img[0]
            if xref in seen:
                continue
            seen.add(xref)
            try:
                pix = fitz.Pixmap(doc, xref)
            except Exception:
                continue
            if pix.width < MIN_FIG_PX or pix.height < MIN_FIG_PX:
                continue
            fid += 1
            fname = "figure_{:02d}.png".format(fid)
            fpath = os.path.join(fig_dir, fname)
            try:
                if pix.n - pix.alpha >= 4:  # CMYK -> RGB
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                pix.save(fpath)
            except Exception:
                fid -= 1
                continue

            rects = page.get_image_rects(xref)
            img_bottom = max((r.y1 for r in rects), default=None)
            caption = _nearest_caption(caps, img_bottom, prefer_above=False) \
                if img_bottom is not None else None
            if caption:
                matched_caption_keys.add((pno, caption[0]))
            figures.append({
                "id": "FIG-{:02d}".format(fid),
                "detected_number": caption[0] if caption else None,
                "caption": caption[1] if caption else None,
                "page": pno,
                "image_file": fname,
                "surrounding_text": _surrounding_text(page, rects),
            })

    # Captions with no matched raster image -> likely vector figures.
    for pno, caps in all_caps.items():
        for num, title, y0 in caps:
            if (pno, num) in matched_caption_keys:
                continue
            unmatched_images.append({
                "detected_number": num,
                "caption": title,
                "page": pno,
                "image_file": None,
                "note": "No embedded raster image matched. Render page {} to "
                        "recover this figure if visual review is needed.".format(pno),
            })

    return {"figures": figures, "unmatched_captions": unmatched_images}


def _nearest_caption(caps, y_ref, prefer_above):
    """Pick the caption line closest to a table/figure edge.

    Table captions usually sit above the table; figure captions below the image.
    """
    if not caps or y_ref is None:
        return None
    best = None
    best_d = 1e9
    for num, title, y0 in caps:
        if prefer_above and y0 > y_ref:
            continue  # caption should be above the table top
        if not prefer_above and y0 < y_ref:
            continue  # caption should be below the image bottom
        d = abs(y0 - y_ref)
        if d < best_d:
            best_d, best = d, (num, title)
    if best is None:  # fall back to plain nearest
        for num, title, y0 in caps:
            d = abs(y0 - y_ref)
            if d < best_d:
                best_d, best = d, (num, title)
    return best


def _surrounding_text(page, rects, margin=60):
    if not rects:
        return ""
    r = rects[0]
    clip = fitz.Rect(0, max(0, r.y0 - margin), page.rect.width,
                     min(page.rect.height, r.y1 + margin))
    return page.get_text("text", clip=clip).strip()[:600]


def render_page(run_id, page_no, dpi=170):
    """On-demand render of a single page to workspace/<run-id>/pages/page_NN.png."""
    manifest = common.load_manifest(run_id)
    src = os.path.join(common.subdir(run_id, "inputs"), manifest["proposal_file"])
    doc = fitz.open(src)
    page = doc[page_no - 1]
    out = os.path.join(common.subdir(run_id, "pages"),
                       "page_{:02d}.png".format(page_no))
    page.get_pixmap(dpi=dpi).save(out)
    doc.close()
    return out


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------
def extract_docx(run_id, path):
    import docx

    d = docx.Document(path)
    headings = []
    text_parts = []
    for p in d.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        text_parts.append(txt)
        style = (p.style.name or "") if p.style else ""
        if style.lower().startswith("heading") or style.lower() == "title":
            headings.append({"text": txt, "style": style})
    full_text = "\n".join(text_parts)

    document = {
        "run_id": run_id, "source": os.path.basename(path), "type": "docx",
        "page_count": None,
        "pages": [{"page": None, "text": full_text, "headings": headings}],
        "full_text": full_text,
    }
    common.write_json(common.internal_path(run_id, "document.json"), document)

    tables = []
    for idx, t in enumerate(d.tables, 1):
        rows = [[c.text.strip() for c in row.cells] for row in t.rows]
        headers = rows[0] if rows else []
        tables.append({
            "id": "TAB-{:02d}".format(idx), "detected_number": None,
            "detected_title": None, "page": None,
            "column_headers": headers, "row_count": len(rows),
            "sample_rows": rows[:3],
        })
    labels = [{"detected_number": int(common.TABLE_CAPTION_RE.match(ln).group(1)),
               "detected_title": common.TABLE_CAPTION_RE.match(ln).group(2).strip(),
               "page": None}
              for ln in text_parts if common.TABLE_CAPTION_RE.match(ln)]
    common.write_json(common.internal_path(run_id, "table-map.json"),
                      {"tables": tables, "detected_table_labels": labels})

    figure_map = _extract_docx_figures(run_id, d, text_parts)
    common.write_json(common.internal_path(run_id, "figure-map.json"), figure_map)
    return document, {"tables": tables}, figure_map


def _extract_docx_figures(run_id, d, text_parts):
    fig_dir = common.subdir(run_id, "figures")
    captions = [(int(common.FIGURE_CAPTION_RE.match(ln).group(1)),
                 common.FIGURE_CAPTION_RE.match(ln).group(2).strip())
                for ln in text_parts if common.FIGURE_CAPTION_RE.match(ln)]
    figures = []
    fid = 0
    for rel in d.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            blob = rel.target_part.blob
        except Exception:
            continue
        ext = os.path.splitext(rel.target_part.partname)[1] or ".png"
        fid += 1
        fname = "figure_{:02d}{}".format(fid, ext)
        with open(os.path.join(fig_dir, fname), "wb") as fh:
            fh.write(blob)
        cap = captions[fid - 1] if fid - 1 < len(captions) else (None, None)
        figures.append({
            "id": "FIG-{:02d}".format(fid), "detected_number": cap[0],
            "caption": cap[1], "page": None, "image_file": fname,
            "surrounding_text": "",
        })
    return {"figures": figures, "unmatched_captions": []}


# ---------------------------------------------------------------------------
def main():
    if len(sys.argv) < 2:
        raise SystemExit("Usage: python3 scripts/extract.py <run-id>")
    run_id = sys.argv[1]
    manifest = common.load_manifest(run_id)
    src = os.path.join(common.subdir(run_id, "inputs"), manifest["proposal_file"])

    if manifest["proposal_type"] == "pdf":
        doc, tables, figures = extract_pdf(run_id, src)
    else:
        doc, tables, figures = extract_docx(run_id, src)

    manifest.setdefault("steps_completed", []).append("extract")
    common.write_json(common.internal_path(run_id, "manifest.json"), manifest)

    n_tables = len(tables.get("tables", []))
    n_figs = len(figures.get("figures", []))
    n_unmatched = len(figures.get("unmatched_captions", []))
    print("Extraction complete for run '{}':".format(run_id))
    print("  pages           : {}".format(doc.get("page_count")))
    print("  tables detected : {}".format(n_tables))
    print("  figures (raster): {}".format(n_figs))
    print("  figure captions without raster image: {}".format(n_unmatched))
    print("  figures dir     : {}".format(common.subdir(run_id, "figures")))


if __name__ == "__main__":
    main()
